import io
import zipfile
import re
import olefile
from typing import Union
from pdfminer.high_level import extract_text
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from docx import Document

def get_file_bytes(file) -> bytes:
    """Конвертирует файл в bytes"""
    if hasattr(file, 'read'):
        file.seek(0)
        return file.read()
    return file

def is_valid_pdf(file_bytes: bytes) -> bool:
    """Проверяет валидность PDF"""
    try:
        with io.BytesIO(file_bytes) as f:
            parser = PDFParser(f)
            PDFDocument(parser)
            return True
    except:
        return False

def parse_pdf(file_bytes: bytes) -> str:
    """Парсинг PDF файлов"""
    try:
        if not is_valid_pdf(file_bytes):
            raise ValueError("Неверный формат PDF или файл поврежден")
        
        text = extract_text(io.BytesIO(file_bytes))
        if not text.strip():
            raise ValueError("PDF не содержит текста (возможно сканированный документ)")
        
        return text.strip()
    except Exception as e:
        raise ValueError(f"PDF ошибка: {str(e)}")

def parse_docx(file_bytes: bytes) -> str:
    """Парсинг DOCX файлов"""
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        
        if not text.strip():
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                with z.open('word/document.xml') as f:
                    xml_content = f.read().decode('utf-8')
                    text = re.sub('<[^>]+>', '', xml_content)[:10000]
        
        return text.strip()
    except Exception as e:
        raise ValueError(f"DOCX ошибка: {str(e)}")

def parse_old_doc(file_bytes: bytes) -> str:
    """Парсинг старых DOC файлов"""
    try:
        with io.BytesIO(file_bytes) as f:
            if not olefile.isOleFile(f):
                raise ValueError("Не является DOC файлом")
            
            ole = olefile.OleFileIO(f)
            if ole.exists('WordDocument'):
                stream = ole.openstream('WordDocument')
                text = stream.read().decode('latin-1', errors='ignore')
                return ' '.join(text.split()).strip()
        
        raise ValueError("Не удалось извлечь текст из DOC")
    except Exception as e:
        raise ValueError(f"DOC ошибка: {str(e)}")

def parse_txt(file_bytes: bytes) -> str:
    """Парсинг текстовых файлов"""
    encodings = ['utf-8', 'cp1251', 'windows-1251']
    for enc in encodings:
        try:
            return file_bytes.decode(enc)
        except UnicodeDecodeError:
            continue
    raise ValueError("Не удалось определить кодировку")

def parse_resume(file) -> str:
    """Основная функция для обработки резюме"""
    try:
        file_bytes = get_file_bytes(file)
        
        if file_bytes.startswith(b'%PDF-') or is_valid_pdf(file_bytes):
            return parse_pdf(file_bytes)
        elif file_bytes.startswith(b'PK\x03\x04'):
            return parse_docx(file_bytes)
        elif file_bytes.startswith(b'\xD0\xCF\x11\xE0'):
            return parse_old_doc(file_bytes)
        else:
            try:
                return parse_txt(file_bytes)
            except:
                raise ValueError("Неподдерживаемый формат файла")
                
    except Exception as e:
        raise ValueError(f"Ошибка обработки файла: {str(e)}")